import PyPDF2
import pdfplumber
import docx
import re
import os

def extract_text_from_pdf(file_path):
    """Extract text from PDF using pdfplumber with PyPDF2 fallback."""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
        if len(text.strip()) > 100:
            return text
    except Exception as e:
        print(f"pdfplumber failed: {e}")

    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() + "\n\n"
    except Exception as e:
        print(f"PyPDF2 failed: {e}")

    return text

def extract_text_from_docx(file_path):
    """Extract text from Word document."""
    try:
        doc = docx.Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " | "
                text += "\n"
        return text
    except Exception as e:
        print(f"docx extraction failed: {e}")
        return ""

def extract_text_from_txt(file_path):
    """Extract text from plain text file."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        print(f"txt extraction failed: {e}")
        return ""

def parse_document(file_path):
    """Parse document based on file type."""
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        text = extract_text_from_docx(file_path)
    elif ext in ['.txt', '.md']:
        text = extract_text_from_txt(file_path)
    else:
        return {"success": False, "error": f"Unsupported file type: {ext}"}
    
    if not text or len(text.strip()) < 50:
        return {"success": False, "error": "Could not extract text from document. Please ensure the document is not scanned/image-based."}
    
    return {
        "success": True,
        "text": text,
        "word_count": len(text.split()),
        "char_count": len(text)
    }

def extract_key_sections(text):
    """Extract key sections from RFP/NOFO document."""
    sections = {
        "funding_opportunity": "",
        "eligibility": "",
        "deadline": "",
        "funding_amount": "",
        "purpose": "",
        "requirements": "",
        "evaluation_criteria": "",
        "submission_requirements": "",
        "contact_info": ""
    }
    
    # Extract deadline
    deadline_patterns = [
        r'deadline[:\s]+([^\n]+)',
        r'due date[:\s]+([^\n]+)',
        r'applications? (?:are )?due[:\s]+([^\n]+)',
        r'submission deadline[:\s]+([^\n]+)',
        r'closing date[:\s]+([^\n]+)'
    ]
    for pattern in deadline_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            sections["deadline"] = match.group(1).strip()
            break
    
    # Extract funding amount
    amount_patterns = [
        r'(?:award|funding|grant) (?:amount|ceiling|maximum|up to)[:\s]+\$?([\d,]+(?:\.\d+)?(?:\s*(?:million|thousand|M|K))?)',
        r'\$\s*([\d,]+(?:\.\d+)?(?:\s*(?:million|thousand|M|K))?)\s*(?:per award|per grant|available)',
        r'total (?:funding|award)[:\s]+\$?([\d,]+(?:\.\d+)?(?:\s*(?:million|thousand|M|K))?)'
    ]
    for pattern in amount_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            sections["funding_amount"] = match.group(0).strip()
            break

    # Extract eligibility
    elig_match = re.search(
        r'(?:eligib(?:ility|le)[^\n]*\n)((?:.*\n){1,15})',
        text, re.IGNORECASE
    )
    if elig_match:
        sections["eligibility"] = elig_match.group(0)[:500]

    # Extract purpose/overview
    purpose_match = re.search(
        r'(?:purpose|overview|background|program description)[:\s]*\n((?:.*\n){1,20})',
        text, re.IGNORECASE
    )
    if purpose_match:
        sections["purpose"] = purpose_match.group(0)[:1000]

    return sections